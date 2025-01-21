from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import json
import os
import requests
from datetime import datetime
#import locale

app = Flask(__name__)

# Load members from JSON file
def load_members():
    if os.path.exists("members.json"):
        with open("members.json", "r") as file:
            return json.load(file)
    return []

# Save members to JSON file
def save_members(members):
    with open("members.json", "w") as file:
        json.dump(members, file, indent=4)

# Set font for a Word document run
def set_font(run, font_name="Arial", font_size=12, font_color=(0, 0, 0), bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*font_color)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

# Generate a Word document
def generate_docx(data, signer, task_details, output_path):
    template = Document("SPT_TEMPLATE.docx")

    for paragraph in template.paragraphs:
        if "{date}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{date}", datetime.now().strftime("%d %B %Y"))

    tables = template.tables

    # Table 1: Signer Details
    header_table = tables[0]
    header_data = [
    signer['name'].title(),
    signer['nip'],
    signer['pangkat'],
    # Check if the signer['jabatan'] matches the required value
    signer['jabatan'] if signer['jabatan'] == "Direktur Seismologi Teknik Geofisika Potensial dan Tanda Waktu" \
    else "Plh. Direktur Seismologi Teknik Geofisika Potensial dan Tanda Waktu" \
    if "Kepala" in signer['jabatan'] or "Deputi" in signer['jabatan'] \
    else signer['jabatan'],
    signer['organization']
]
    for i, value in enumerate(header_data):
        cell = header_table.cell(i, 2)
        cell.text = value
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=(i == 0))  # Make the first value bold

    # Table 2: Assignment Details (Multiple Members)
    assignments_table = tables[1]
    current_row_idx = 0
    field_names_translated = ["Nama", "NIP", "Pangkat/Golongan", "Jabatan", "Satuan Organisasi"]
    for row_idx, member in enumerate(data):
        for field_idx, field_name in enumerate(["name", "nip", "pangkat", "jabatan", "organization"]):
            if current_row_idx < len(assignments_table.rows):
                row = assignments_table.rows[current_row_idx].cells
            else:
                row = assignments_table.add_row().cells
            
            row[0].text = str(row_idx + 1) if field_idx == 0 else ""
            row[1].text = field_names_translated[field_idx]
            row[2].text = ":"
            row[3].text = member[field_name]

            if field_name == "name":
                row[3].text = member[field_name].title()  # Capitalize each word for names
            else:
                row[3].text = member[field_name]
            
            
            paragraph1 = row[0].paragraphs[0]
            paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph2 = row[2].paragraphs[0]
            paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for cell in row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = Pt(12)  # Line spacing
                    paragraph.paragraph_format.space_before = Pt(0)  # Space before paragraph
                    paragraph.paragraph_format.space_after = Pt(0)  # Space after paragraph
                for run in paragraph.runs:
                    is_name_field = field_name == "name" and cell is row[3]
                    set_font(run, bold=is_name_field)

            current_row_idx += 1
        

        if row_idx < len(data) - 1:
            if current_row_idx < len(assignments_table.rows):
                row = assignments_table.rows[current_row_idx].cells
            else:
                row = assignments_table.add_row().cells
            for cell in row:
                cell.text = ""
            current_row_idx += 1

    # Table 3: Task Details
    
    task_table = tables[2]
    task_table.cell(0, 2).text = task_details["tugas"]
    task_table.cell(1, 2).text = task_details["lama_perjalanan"]
    task_table.cell(2, 2).text = task_details["lokasi"]
    #task_table.cell(3, 2).text = formatted_date
    task_table.cell(3, 2).text = task_details["tanggal_berangkat"] 
    task_table.cell(4, 2).text = task_details["sumber_dana"]

    for row in task_table.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run)

    #Table 4: Signed
    signed_cell = tables[3]
    current_month_year = datetime.now().strftime("%B %Y")
    signed_cell.cell(0, 0).text = f"Jakarta,     {current_month_year}"
    signed_cell.cell(1, 0).text = signer['jabatan']
    # Check if the signer['jabatan'] matches the required value
    signer['jabatan'] if signer['jabatan'] == "Direktur Seismologi Teknik Geofisika Potensial dan Tanda Waktu" else "Plh. Direktur Seismologi Teknik Geofisika Potensial dan Tanda Waktu"
    signed_cell.cell(3, 0).text = signer['name'].title()

    for row in signed_cell.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run)


    template.save(output_path)

# Flask routes
@app.route('/')
def index():
    members = load_members()
    return render_template('index.html', members=members)

@app.route('/search_member', methods=['GET'])
def search_member():
    query = request.args.get('query', '').lower()

    # Fetch data from the API
    try:
        response = requests.get("http://202.90.198.220/api/sdm/sdm-csv.bmkg")
        response.raise_for_status()  # Raise an exception for HTTP errors
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to fetch data from the API: {e}"}), 500

    # Process the API response
    data = response.text.splitlines()
    results = []

    # Parse the CSV data and filter by the search query
    for line in data:
        fields = line.split(";")
        if len(fields) >= 6 and query in fields[1].lower():  # Check if the query matches the name
            if fields[3] == "I/a":
                fields[3] = "Juru Muda / Ia"
            elif fields[3] == "I/b":
                fields[3] = "Juru Muda Tk. I / Ib"
            elif fields[3] == "I/c":
                fields[3] = "Juru / Ic"
            elif fields[3] == "I/d":
                fields[3] = "Juru / Id"
            elif fields[3] == "II/a":
                fields[3] = "Pengatur Muda / IIa"
            elif fields[3] == "II/b":
                fields[3] = "Pengatur Muda Tk. I / IIb"
            elif fields[3] == "II/c":
                fields[3] = "Pengatur / IIc"
            elif fields[3] == "II/d":
                fields[3] = "Pengatur Tingkat I / IId"
            elif fields[3] == "III/a":
                fields[3] = "Penata Muda / IIIa"
            elif fields[3] == "III/b":
                fields[3] = "Penata Muda Tk. I / IIIb"
            elif fields[3] == "III/c":
                fields[3] = "Penata / IIIc"
            elif fields[3] == "III/d":
                fields[3] = "Penata Tk. I / IIId"
            elif fields[3] == "IV/a":
                fields[3] = "Pembina / IVa"
            elif fields[3] == "IV/b":
                fields[3] = "Pembina Tk. I / IVb"
            elif fields[3] == "IV/c":
                fields[3] = "Pembina Muda / IVc"
            elif fields[3] == "IV/d":
                fields[3] = "Pembina Madya / IVd"
            elif fields[3] == "IV/e":
                fields[3] = "Pembina Utama / IVe"
            else : fields[3]
            results.append({
                "nip": fields[0].strip(),
                "name": fields[1].strip(),
                "pangkat": fields[3].strip(),
                "jabatan": fields[4].strip(),
                "organization": fields[5].strip()
            })

    return jsonify(results)

@app.route('/generate_st', methods=['POST'])
def generate_document():
    try:
        data = request.get_json()  # Use get_json() instead of request.json
        
        if not data:
            return jsonify({"error": "No JSON data received"}), 400
            
        # Extract data
        members = data.get('members', [])
        signer = data.get('signer', {})
        task_details = data.get('task_details', {})

        if not members or not signer or not task_details:
            return jsonify({"error": "Missing required data"}), 400

        # Generate document file path
        output_path = f"Surat_Tugas_{task_details['tugas']}_{task_details['tanggal_berangkat']}.docx"
        
        # Generate the Word document
        generate_docx(members, signer, task_details, output_path)
        send_file(output_path, as_attachment=True)
        
        # Return success response
        return jsonify({"status": "success", "filename": output_path}), 200
        
    except Exception as e:
        app.logger.error(f"Error in generate_document: {str(e)}")
        return jsonify({"error": f"Error generating document: {str(e)}"}), 500

@app.route('/download_st', methods=['POST'])
def download_document():
    data = request.get_json()  # Use get_json() instead of request.json
       
    # Generate document file path
    task_details = data.get('task_details', {})
    output_path = f"Surat_Tugas_{task_details['tugas']}_{task_details['tanggal_berangkat']}.docx"
        
      
    return send_file(output_path, as_attachment=True)

@app.route('/add_member', methods=['POST'])
def add_member():
    new_member = request.json
    members = load_members()
    members.append(new_member)
    save_members(members)
    return jsonify({"message": "Member added successfully!", "member": new_member})


if __name__ == '__main__':
    app.run(debug=True)
